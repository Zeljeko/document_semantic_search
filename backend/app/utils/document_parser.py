import os
import PyPDF2 
from docx import Document
from typing import List, Tuple
import re

class DocumentParser:
    """
    Handles parsing of PDF, DOCX, TXT files.
    Return clean text and metadata.
    """

    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, dict]:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                    
                metadata = {
                    "num_of_pages": len(pdf_reader.pages),
                    "file_type": "pdf"
                }
                return text, metadata
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, dict]:
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "file_type": "docx"
            }

            return text, metadata
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        
    @staticmethod
    def parse_txt(file_path:str) -> Tuple[str, dict]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
                metadata = {
                    "num_characters": len(text),
                    "file_type": "txt"
                }

                return text, metadata
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
        
    @classmethod
    def parse_document(cls, file_path:str) -> Tuple[str, dict]:
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return cls.parse_pdf(file_path)
        elif file_extension == '.docx':
            return cls.parse_docx(file_path)
        elif file_extension == '.txt':
            return cls.parse_txt(file_path)
        else:
            return ValueError(f"Unsupported File Type: {file_extension}")

